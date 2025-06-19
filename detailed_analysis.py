#!/usr/bin/env python3
"""
Detailed Student Feedback Analysis for Specific Courses
Focuses on the three main courses mentioned in the request
"""

import os
import re
from collections import defaultdict
from docx import Document
import json
from pathlib import Path

class DetailedFeedbackAnalyzer:
    def __init__(self, data_path):
        self.data_path = Path(data_path)
        self.detailed_results = {}
        
    def read_docx_safely(self, file_path):
        """Safely read a docx file and return its content"""
        try:
            doc = Document(file_path)
            content = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content.append(text)
            
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            content.append(text)
            
            return '\n'.join(content)
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return ""
    
    def analyze_saturday_primary_course(self):
        """Analyze Saturday - 1.5 - 3 PM - 02IPDEC2404 - PSD I"""
        course_path = self.data_path / 'primary' / 'Saturday - 1.5 - 3 PM - 02IPDEC2404 - PSD I'
        
        if not course_path.exists():
            print(f"Course path not found: {course_path}")
            return
        
        print(f"\\nAnalyzing: {course_path.name}")
        
        course_data = {
            'course_code': '02IPDEC2404',
            'schedule': 'Saturday 1:30-3:00 PM',
            'level': 'Primary PSD I',
            'students': {},
            'units_analyzed': []
        }
        
        # Get all unit folders
        unit_folders = [f for f in course_path.iterdir() if f.is_dir()]
        unit_folders.sort(key=lambda x: x.name)
        
        for unit_folder in unit_folders:
            unit_name = unit_folder.name
            print(f"  Processing unit: {unit_name}")
            
            # Look for individual student files
            student_files = [f for f in unit_folder.iterdir() if f.suffix == '.docx']
            
            for student_file in student_files:
                filename = student_file.stem
                
                # Extract student name from filename
                student_name = None
                if ' - Unit ' in filename:
                    student_name = filename.split(' - Unit ')[0].strip()
                elif ' - ' in filename and 'Feedback' in filename:
                    student_name = filename.split(' - ')[0].strip()
                elif filename.endswith('.docx'):
                    # Check if it's not a generic unit file
                    if not filename.startswith(('Unit ', 'Copy of', '1.', '2.', '3.', '4.')):
                        student_name = filename.strip()
                
                if student_name and student_name not in ['Copy of', 'Unit']:
                    content = self.read_docx_safely(student_file)
                    
                    if student_name not in course_data['students']:
                        course_data['students'][student_name] = {'units': {}, 'progression': []}
                    
                    course_data['students'][student_name]['units'][unit_name] = {
                        'content': content,
                        'file_path': str(student_file),
                        'content_length': len(content)
                    }
            
            # Also check for general unit feedback files
            general_files = [f for f in unit_folder.iterdir() if f.suffix == '.docx' and 
                           (f.stem.startswith(unit_name) or f.stem.startswith(('1.', '2.', '3.', '4.')))]
            
            for general_file in general_files:
                content = self.read_docx_safely(general_file)
                if content and len(content) > 100:  # Only process substantial content
                    course_data['units_analyzed'].append({
                        'unit': unit_name,
                        'file': general_file.name,
                        'content_preview': content[:200] + '...' if len(content) > 200 else content
                    })
        
        self.detailed_results['saturday_primary'] = course_data
        return course_data
    
    def analyze_thursday_primary_course(self):
        """Analyze Thursday - 6 - 7.5 - 02IPDEC2401 - PSD I"""
        course_path = self.data_path / 'primary' / 'Thursday - 6 - 7.5 - 02IPDEC2401 - PSD I'
        
        if not course_path.exists():
            print(f"Course path not found: {course_path}")
            return
        
        print(f"\\nAnalyzing: {course_path.name}")
        
        course_data = {
            'course_code': '02IPDEC2401',
            'schedule': 'Thursday 6:00-7:30 PM',
            'level': 'Primary PSD I',
            'students': {},
            'units_analyzed': []
        }
        
        # Special handling for 1.1 folder with individual student files
        unit_1_1_path = course_path / '1.1'
        if unit_1_1_path.exists():
            print(f"  Processing unit: 1.1 (individual files)")
            
            student_files = [f for f in unit_1_1_path.iterdir() if f.suffix == '.docx']
            for student_file in student_files:
                filename = student_file.stem
                student_name = filename.replace(' - Primary Feedback Sheet', '').strip()
                
                if student_name:
                    content = self.read_docx_safely(student_file)
                    
                    if student_name not in course_data['students']:
                        course_data['students'][student_name] = {'units': {}, 'progression': []}
                    
                    course_data['students'][student_name]['units']['1.1'] = {
                        'content': content,
                        'file_path': str(student_file),
                        'content_length': len(content)
                    }
        
        # Process other units
        unit_folders = [f for f in course_path.iterdir() if f.is_dir() and f.name != '1.1']
        unit_folders.sort(key=lambda x: x.name)
        
        for unit_folder in unit_folders:
            unit_name = unit_folder.name
            print(f"  Processing unit: {unit_name}")
            
            # Look for unit feedback files
            unit_files = [f for f in unit_folder.iterdir() if f.suffix == '.docx']
            
            for unit_file in unit_files:
                content = self.read_docx_safely(unit_file)
                if content and len(content) > 100:
                    course_data['units_analyzed'].append({
                        'unit': unit_name,
                        'file': unit_file.name,
                        'content_preview': content[:200] + '...' if len(content) > 200 else content
                    })
        
        self.detailed_results['thursday_primary'] = course_data
        return course_data
    
    def analyze_saturday_secondary_course(self):
        """Analyze Saturday - 3_00 - 4_30 - 01IPDED2404 - PSD I"""
        course_path = self.data_path / 'secondary' / 'Saturday - 3_00 - 4_30 -  01IPDED2404 - PSD I'
        
        if not course_path.exists():
            print(f"Course path not found: {course_path}")
            return
        
        print(f"\\nAnalyzing: {course_path.name}")
        
        course_data = {
            'course_code': '01IPDED2404',
            'schedule': 'Saturday 3:00-4:30 PM',
            'level': 'Secondary PSD I',
            'students': {},
            'units_analyzed': []
        }
        
        # Special handling for 1.1 folder with individual student files
        unit_1_1_path = course_path / '1.1'
        if unit_1_1_path.exists():
            print(f"  Processing unit: 1.1 (individual files)")
            
            student_files = [f for f in unit_1_1_path.iterdir() if f.suffix == '.docx' and 
                           ' - Unit 1.1 Feedback' in f.stem]
            
            for student_file in student_files:
                filename = student_file.stem
                student_name = filename.replace(' - Unit 1.1 Feedback', '').strip()
                
                if student_name:
                    content = self.read_docx_safely(student_file)
                    
                    if student_name not in course_data['students']:
                        course_data['students'][student_name] = {'units': {}, 'progression': []}
                    
                    course_data['students'][student_name]['units']['1.1'] = {
                        'content': content,
                        'file_path': str(student_file),
                        'content_length': len(content)
                    }
        
        # Process other unit files (they appear to be direct files, not in folders)
        unit_files = [f for f in course_path.iterdir() if f.suffix == '.docx' and 
                     not f.name.startswith('Copy of')]
        
        for unit_file in unit_files:
            filename = unit_file.stem
            
            # Extract unit info from filename
            unit_match = re.search(r'(\\d+\\.\\d+)', filename)
            if unit_match:
                unit_name = unit_match.group(1)
                content = self.read_docx_safely(unit_file)
                
                if content and len(content) > 100:
                    course_data['units_analyzed'].append({
                        'unit': unit_name,
                        'file': unit_file.name,
                        'content_preview': content[:300] + '...' if len(content) > 300 else content
                    })
        
        # Process Unit folders
        unit_folders = [f for f in course_path.iterdir() if f.is_dir() and f.name.startswith('Unit')]
        unit_folders.sort(key=lambda x: x.name)
        
        for unit_folder in unit_folders:
            unit_name = unit_folder.name
            print(f"  Processing {unit_name}")
            
            unit_files = [f for f in unit_folder.iterdir() if f.suffix == '.docx']
            
            for unit_file in unit_files:
                content = self.read_docx_safely(unit_file)
                if content and len(content) > 100:
                    course_data['units_analyzed'].append({
                        'unit': unit_name,
                        'file': unit_file.name,
                        'content_preview': content[:300] + '...' if len(content) > 300 else content
                    })
        
        self.detailed_results['saturday_secondary'] = course_data
        return course_data
    
    def find_student_progression(self, student_name, course_data):
        """Track a specific student's progression across units"""
        if student_name not in course_data['students']:
            return None
        
        student_data = course_data['students'][student_name]
        units = student_data['units']
        
        progression = []
        for unit, data in sorted(units.items()):
            analysis = {
                'unit': unit,
                'content_length': data['content_length'],
                'content_preview': data['content'][:200] + '...' if len(data['content']) > 200 else data['content'],
                'themes': self.extract_themes(data['content']),
                'feedback_highlights': self.extract_feedback_highlights(data['content'])
            }
            progression.append(analysis)
        
        return progression
    
    def extract_themes(self, content):
        """Extract common themes from feedback content"""
        content_lower = content.lower()
        themes = []
        
        theme_keywords = {
            'confidence': ['confidence', 'confident', 'self-assured'],
            'participation': ['participation', 'participate', 'engaged', 'involvement'],
            'improvement': ['improvement', 'improved', 'better', 'progress', 'growth'],
            'challenges': ['challenge', 'difficult', 'struggle', 'need to work on'],
            'speaking_skills': ['speaking', 'voice', 'volume', 'clarity', 'pronunciation'],
            'listening': ['listening', 'attention', 'focus'],
            'teamwork': ['teamwork', 'collaboration', 'group work', 'cooperation'],
            'creativity': ['creative', 'imagination', 'original', 'innovative']
        }
        
        for theme, keywords in theme_keywords.items():
            if any(keyword in content_lower for keyword in keywords):
                themes.append(theme)
        
        return themes
    
    def extract_feedback_highlights(self, content):
        """Extract key feedback points from content"""
        if not content or len(content) < 50:
            return []
        
        # Split into sentences and take meaningful ones
        sentences = [s.strip() for s in content.split('.') if len(s.strip()) > 20]
        
        # Return first few substantial sentences
        return sentences[:3]
    
    def generate_comprehensive_report(self):
        """Generate the comprehensive report requested"""
        print("\\n" + "="*100)
        print("COMPREHENSIVE STUDENT FEEDBACK ANALYSIS REPORT")
        print("="*100)
        
        # Run all analyses
        saturday_primary = self.analyze_saturday_primary_course()
        thursday_primary = self.analyze_thursday_primary_course()
        saturday_secondary = self.analyze_saturday_secondary_course()
        
        # Consolidate all students
        all_students = set()
        course_mappings = {}
        
        for course_key, course_data in self.detailed_results.items():
            if course_data:
                for student in course_data['students'].keys():
                    all_students.add(student)
                    course_mappings[student] = {
                        'course_code': course_data['course_code'],
                        'schedule': course_data['schedule'],
                        'level': course_data['level']
                    }
        
        print(f"\\n1. STUDENT NAME EXTRACTION")
        print(f"   Total unique students found: {len(all_students)}")
        print(f"   Students by course:")
        
        for course_key, course_data in self.detailed_results.items():
            if course_data:
                print(f"\\n   {course_data['course_code']} - {course_data['schedule']}:")
                print(f"     Level: {course_data['level']}")
                students = list(course_data['students'].keys())
                print(f"     Students ({len(students)}): {', '.join(sorted(students))}")
        
        print(f"\\n2. COURSE CODE ANALYSIS")
        course_codes = set()
        for course_data in self.detailed_results.values():
            if course_data:
                course_codes.add(course_data['course_code'])
        
        print(f"   Course codes identified: {sorted(course_codes)}")
        
        print(f"\\n3. STUDENT PROGRESSION TRACKING")
        
        # Find students with multiple units for detailed progression analysis
        students_with_progression = []
        for course_key, course_data in self.detailed_results.items():
            if course_data:
                for student_name, student_data in course_data['students'].items():
                    if len(student_data['units']) > 1:
                        students_with_progression.append((student_name, course_data, len(student_data['units'])))
        
        # Sort by number of units (descending)
        students_with_progression.sort(key=lambda x: x[2], reverse=True)
        
        # Analyze top 3 students with most progression data
        for i, (student_name, course_data, unit_count) in enumerate(students_with_progression[:3]):
            print(f"\\n   Student {i+1}: {student_name}")
            print(f"   Course: {course_data['course_code']} - {course_data['schedule']}")
            print(f"   Units with feedback: {unit_count}")
            
            progression = self.find_student_progression(student_name, course_data)
            if progression:
                print(f"   Progression analysis:")
                for prog in progression:
                    print(f"     Unit {prog['unit']}:")
                    print(f"       Content length: {prog['content_length']} characters")
                    if prog['themes']:
                        print(f"       Themes: {', '.join(prog['themes'])}")
                    if prog['feedback_highlights']:
                        print(f"       Key feedback: {prog['feedback_highlights'][0][:100]}...")
        
        print(f"\\n4. SAMPLE DEEP ANALYSIS")
        
        # Show sample content from different courses
        for course_key, course_data in self.detailed_results.items():
            if course_data and course_data['units_analyzed']:
                print(f"\\n   {course_data['course_code']} - Sample Unit Content:")
                for unit in course_data['units_analyzed'][:2]:  # Show first 2 units
                    print(f"     {unit['unit']} ({unit['file']}):")
                    print(f"       {unit['content_preview']}")
        
        # Save detailed results
        output_file = "/Users/tikaram/Downloads/Claude Code/student-growth/growth-compass/detailed_feedback_analysis.json"
        with open(output_file, 'w') as f:
            json.dump(self.detailed_results, f, indent=2)
        
        print(f"\\nDetailed analysis saved to: {output_file}")
        
        return {
            'total_students': len(all_students),
            'course_codes': sorted(course_codes),
            'students_with_progression': len(students_with_progression),
            'courses_analyzed': len(self.detailed_results)
        }

def main():
    data_path = "/Users/tikaram/Downloads/Claude Code/student-growth/growth-compass/data"
    
    analyzer = DetailedFeedbackAnalyzer(data_path)
    results = analyzer.generate_comprehensive_report()
    
    print(f"\\nAnalysis Summary:")
    print(f"  Total students: {results['total_students']}")
    print(f"  Course codes: {results['course_codes']}")
    print(f"  Students with progression data: {results['students_with_progression']}")
    print(f"  Courses analyzed: {results['courses_analyzed']}")

if __name__ == "__main__":
    main()
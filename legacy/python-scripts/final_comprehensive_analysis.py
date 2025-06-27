#!/usr/bin/env python3
"""
Final Comprehensive Student Feedback Analysis
Extracts and analyzes student feedback data from all formats
"""

import os
import re
from collections import defaultdict, Counter
from docx import Document
import json
from pathlib import Path

class ComprehensiveFeedbackAnalyzer:
    def __init__(self, data_path):
        self.data_path = Path(data_path)
        self.all_students = {}  # student_name -> {courses: [], progression: []}
        self.course_info = {}   # course_code -> {schedule, level, students}
        self.detailed_feedback = defaultdict(list)  # student_course_key -> [feedback entries]
        self.student_progression = defaultdict(list)  # For tracking individual growth
        
    def read_docx_with_tables(self, file_path):
        """Read docx file and extract both paragraphs and table content"""
        try:
            doc = Document(file_path)
            content = {
                'paragraphs': [],
                'tables': [],
                'student_name': None,
                'topic': None,
                'feedback': []
            }
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content['paragraphs'].append(paragraph.text.strip())
            
            # Extract table content with structure
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                            
                            # Extract student name
                            if cell_text.startswith('Student Name:') or cell_text.startswith('Student:'):
                                name = cell_text.replace('Student Name:', '').replace('Student:', '').strip()
                                if name:
                                    content['student_name'] = name
                            
                            # Extract topic/motion
                            if 'Topic:' in cell_text or 'Motion:' in cell_text:
                                topic = re.sub(r'^(Topic:|Motion:)', '', cell_text).strip()
                                if topic:
                                    content['topic'] = topic
                            
                            # Extract feedback comments
                            if ('BEST thing' in cell_text or 'NEEDS IMPROVEMENT' in cell_text or 
                                'Teacher comments:' in cell_text or len(cell_text) > 50):
                                content['feedback'].append(cell_text)
                    
                    if row_data:
                        table_data.append(row_data)
                
                if table_data:
                    content['tables'].append(table_data)
            
            return content
            
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return None
    
    def extract_course_info(self, folder_path):
        """Extract course information from folder name"""
        folder_name = folder_path.name
        
        # Different patterns for course folder names
        patterns = [
            r'(\w+)\s*-\s*([\d\s\.\-:_]+)\s*-\s*([A-Z0-9]+)\s*-\s*(.*)',  # Standard format
            r'(\w+)\s*-\s*([\d\s\.\-:_]+)\s*PM\s*-\s*([A-Z0-9\-]+)\s*-\s*(.*)',  # With PM
            r'(\w+)\s*-\s*([\d\s\.\-:_]+)\s*-\s*([A-Z0-9\-]+)\s*/?\s*(.*)',  # Flexible ending
        ]
        
        for pattern in patterns:
            match = re.search(pattern, folder_name)
            if match:
                return {
                    'day': match.group(1).strip(),
                    'time': match.group(2).strip(),
                    'course_code': match.group(3).strip(),
                    'level': match.group(4).strip(),
                    'full_schedule': f"{match.group(1).strip()} {match.group(2).strip()}"
                }
        
        return None
    
    def analyze_course_folder(self, course_folder_path, level_type):
        """Analyze a complete course folder"""
        course_info = self.extract_course_info(course_folder_path)
        if not course_info:
            print(f"Could not parse course info from: {course_folder_path.name}")
            return
        
        print(f"\\nAnalyzing {level_type} course: {course_info['course_code']} - {course_info['full_schedule']}")
        
        # Initialize course info
        course_key = course_info['course_code']
        self.course_info[course_key] = {
            'schedule': course_info['full_schedule'],
            'level': f"{level_type.title()} {course_info['level']}",
            'students': set(),
            'units': []
        }
        
        # Find all unit folders and files
        items = list(course_folder_path.iterdir())
        unit_folders = [item for item in items if item.is_dir()]
        
        for unit_folder in sorted(unit_folders):
            unit_name = unit_folder.name
            print(f"  Processing unit: {unit_name}")
            
            # Look for student files in this unit
            student_files = [f for f in unit_folder.iterdir() if f.suffix == '.docx']
            
            for student_file in student_files:
                self.process_student_file(student_file, course_info, unit_name, level_type)
        
        # Also check for direct unit files (not in folders)
        direct_files = [item for item in items if item.suffix == '.docx']
        for direct_file in direct_files:
            # Try to extract unit info from filename
            unit_match = re.search(r'(\\d+\\.\\d+)', direct_file.stem)
            unit_name = unit_match.group(1) if unit_match else 'unknown'
            self.process_student_file(direct_file, course_info, unit_name, level_type)
    
    def process_student_file(self, file_path, course_info, unit_name, level_type):
        """Process an individual student feedback file"""
        filename = file_path.stem
        
        # Extract student name from filename
        student_name = None
        
        # Various filename patterns
        if ' - Unit ' in filename:
            student_name = filename.split(' - Unit ')[0].strip()
        elif ' - Primary Feedback' in filename:
            student_name = filename.split(' - Primary Feedback')[0].strip()
        elif ' Feedback_' in filename:
            student_name = filename.split(' Feedback_')[0].strip()
        elif ' - ' in filename and not filename.startswith(('1.', '2.', '3.', '4.', 'Unit', 'Copy')):
            parts = filename.split(' - ')
            if len(parts) >= 2 and not parts[0].replace('.', '').isdigit():
                student_name = parts[0].strip()
        elif not filename.startswith(('1.', '2.', '3.', '4.', 'Unit', 'Copy')) and not filename.replace('.', '').isdigit():
            student_name = filename.strip()
        
        # Read file content
        content = self.read_docx_with_tables(file_path)
        if not content:
            return
        
        # Use student name from file content if not extracted from filename
        if not student_name and content['student_name']:
            student_name = content['student_name']
        
        # Clean up student name
        if student_name:
            student_name = re.sub(r'\\s*(Feedback|Sheet)\\s*$', '', student_name).strip()
            student_name = student_name.replace('(', '').replace(')', '').strip()
            
            # Skip non-student names
            if (student_name in ['Copy of', 'Unit', 'Clearing Class', 'Clearing class'] or 
                student_name.startswith('Gabi') or 
                len(student_name) < 2):
                return
        
        if student_name:
            # Add to course students
            course_key = course_info['course_code']
            self.course_info[course_key]['students'].add(student_name)
            
            # Initialize student if not exists
            if student_name not in self.all_students:
                self.all_students[student_name] = {
                    'courses': [],
                    'progression': {}
                }
            
            # Add course info to student
            course_entry = {
                'course_code': course_key,
                'schedule': course_info['full_schedule'],
                'level': f"{level_type.title()} {course_info['level']}"
            }
            if course_entry not in self.all_students[student_name]['courses']:
                self.all_students[student_name]['courses'].append(course_entry)
            
            # Store feedback details
            student_course_key = f"{student_name}_{course_key}"
            feedback_entry = {
                'unit': unit_name,
                'topic': content['topic'],
                'feedback_text': content['feedback'],
                'file_path': str(file_path),
                'content_length': sum(len(f) for f in content['feedback'])
            }
            
            self.detailed_feedback[student_course_key].append(feedback_entry)
            
            # Add to progression tracking
            if course_key not in self.all_students[student_name]['progression']:
                self.all_students[student_name]['progression'][course_key] = []
            
            self.all_students[student_name]['progression'][course_key].append(feedback_entry)
    
    def find_students_with_progression(self, min_units=2):
        """Find students with feedback across multiple units"""
        students_with_progression = []
        
        for student_name, student_data in self.all_students.items():
            for course_code, progression in student_data['progression'].items():
                if len(progression) >= min_units:
                    students_with_progression.append({
                        'name': student_name,
                        'course_code': course_code,
                        'unit_count': len(progression),
                        'progression': sorted(progression, key=lambda x: x['unit'])
                    })
        
        return sorted(students_with_progression, key=lambda x: x['unit_count'], reverse=True)
    
    def analyze_student_growth(self, student_progression):
        """Analyze growth patterns in a student's feedback"""
        analysis = {
            'student': student_progression['name'],
            'course': student_progression['course_code'],
            'total_units': student_progression['unit_count'],
            'growth_indicators': [],
            'persistent_challenges': [],
            'feedback_evolution': []
        }
        
        all_feedback_text = []
        for entry in student_progression['progression']:
            unit_feedback = ' '.join(entry['feedback_text']).lower()
            all_feedback_text.append(unit_feedback)
            
            # Extract key themes from this unit
            unit_analysis = {
                'unit': entry['unit'],
                'topic': entry['topic'],
                'positive_feedback': [],
                'areas_for_improvement': [],
                'specific_skills': []
            }
            
            # Categorize feedback
            for feedback in entry['feedback_text']:
                feedback_lower = feedback.lower()
                
                if any(word in feedback_lower for word in ['great', 'good', 'excellent', 'well done', 'nice']):
                    unit_analysis['positive_feedback'].append(feedback[:100] + '...' if len(feedback) > 100 else feedback)
                
                if any(word in feedback_lower for word in ['needs', 'try to', 'should', 'improve', 'work on']):
                    unit_analysis['areas_for_improvement'].append(feedback[:100] + '...' if len(feedback) > 100 else feedback)
                
                # Look for specific skills mentioned
                skills = ['volume', 'clarity', 'argument', 'rebuttal', 'evidence', 'signposting', 'hook']
                for skill in skills:
                    if skill in feedback_lower:
                        unit_analysis['specific_skills'].append(skill)
            
            analysis['feedback_evolution'].append(unit_analysis)
        
        return analysis
    
    def run_full_analysis(self):
        """Run the complete analysis"""
        print("Starting comprehensive student feedback analysis...")
        print("="*80)
        
        # Analyze primary courses
        primary_path = self.data_path / 'primary'
        if primary_path.exists():
            for course_folder in primary_path.iterdir():
                if course_folder.is_dir() and ' - ' in course_folder.name:
                    self.analyze_course_folder(course_folder, 'primary')
        
        # Analyze secondary courses
        secondary_path = self.data_path / 'secondary'
        if secondary_path.exists():
            for course_folder in secondary_path.iterdir():
                if course_folder.is_dir() and ' - ' in course_folder.name:
                    self.analyze_course_folder(course_folder, 'secondary')
    
    def generate_comprehensive_report(self):
        """Generate the final comprehensive report"""
        print("\\n" + "="*100)
        print("FINAL COMPREHENSIVE STUDENT FEEDBACK ANALYSIS REPORT")
        print("="*100)
        
        # 1. Student Names Extraction
        print("\\n1. STUDENT NAMES EXTRACTED FROM FEEDBACK DOCUMENTS")
        print("-" * 60)
        print(f"Total unique students found: {len(self.all_students)}")
        
        all_student_names = sorted(self.all_students.keys())
        print("\\nComplete student list:")
        for i, name in enumerate(all_student_names, 1):
            print(f"  {i:2d}. {name}")
        
        # 2. Course Code Analysis
        print("\\n2. COURSE CODE ANALYSIS")
        print("-" * 60)
        print(f"Total courses analyzed: {len(self.course_info)}")
        
        for course_code, info in sorted(self.course_info.items()):
            print(f"\\n  Course: {course_code}")
            print(f"    Schedule: {info['schedule']}")
            print(f"    Level: {info['level']}")
            print(f"    Students: {len(info['students'])} - {', '.join(sorted(info['students']))}")
        
        # 3. Student-Course Mapping
        print("\\n3. STUDENT-COURSE MAPPING")
        print("-" * 60)
        
        for student_name in sorted(self.all_students.keys()):
            student_data = self.all_students[student_name]
            print(f"\\n  {student_name}:")
            for course in student_data['courses']:
                units_count = len(student_data['progression'].get(course['course_code'], []))
                print(f"    - {course['course_code']} ({course['schedule']}) - {units_count} units")
        
        # 4. Individual Student Progression
        print("\\n4. INDIVIDUAL STUDENT PROGRESSION TRACKING")
        print("-" * 60)
        
        students_with_progression = self.find_students_with_progression(min_units=2)
        print(f"Students with progression data (2+ units): {len(students_with_progression)}")
        
        # 5. Sample Deep Analysis
        print("\\n5. SAMPLE DEEP ANALYSIS - TOP 3 STUDENTS")
        print("-" * 60)
        
        for i, student_prog in enumerate(students_with_progression[:3]):
            print(f"\\n  STUDENT {i+1}: {student_prog['name']}")
            print(f"  Course: {student_prog['course_code']}")
            print(f"  Units analyzed: {student_prog['unit_count']}")
            
            analysis = self.analyze_student_growth(student_prog)
            
            print("  \\n  Progression Timeline:")
            for j, unit_analysis in enumerate(analysis['feedback_evolution'][:4]):  # Show first 4 units
                print(f"    Unit {unit_analysis['unit']}:")
                print(f"      Topic: {unit_analysis['topic']}")
                if unit_analysis['positive_feedback']:
                    print(f"      Strengths: {unit_analysis['positive_feedback'][0]}")
                if unit_analysis['areas_for_improvement']:
                    print(f"      Growth areas: {unit_analysis['areas_for_improvement'][0]}")
                if unit_analysis['specific_skills']:
                    print(f"      Skills focus: {', '.join(unit_analysis['specific_skills'])}")
        
        # 6. Statistical Summary
        print("\\n6. STATISTICAL SUMMARY")
        print("-" * 60)
        
        # Students by first letter
        letter_count = Counter(name[0] for name in self.all_students.keys())
        print("\\n  Students by first letter:")
        for letter, count in sorted(letter_count.items()):
            print(f"    {letter}: {count}")
        
        # Progression statistics
        all_progression_counts = []
        for student_data in self.all_students.values():
            for course_progression in student_data['progression'].values():
                all_progression_counts.append(len(course_progression))
        
        if all_progression_counts:
            print(f"\\n  Feedback progression statistics:")
            print(f"    Students with feedback: {len(all_progression_counts)}")
            print(f"    Average units per student: {sum(all_progression_counts)/len(all_progression_counts):.1f}")
            print(f"    Maximum units for one student: {max(all_progression_counts)}")
            print(f"    Students with 3+ units: {sum(1 for x in all_progression_counts if x >= 3)}")
        
        # Save detailed results
        output_file = "/Users/tikaram/Downloads/Claude Code/student-growth/growth-compass/final_analysis_results.json"
        
        # Prepare data for JSON serialization
        json_data = {
            'students': {name: {
                'courses': data['courses'],
                'unit_counts': {course: len(progression) for course, progression in data['progression'].items()}
            } for name, data in self.all_students.items()},
            'course_info': {code: {
                'schedule': info['schedule'],
                'level': info['level'],
                'students': list(info['students'])
            } for code, info in self.course_info.items()},
            'students_with_progression': [{
                'name': sp['name'],
                'course_code': sp['course_code'],
                'unit_count': sp['unit_count']
            } for sp in students_with_progression]
        }
        
        with open(output_file, 'w') as f:
            json.dump(json_data, f, indent=2)
        
        print(f"\\nDetailed results saved to: {output_file}")
        
        return {
            'total_students': len(self.all_students),
            'total_courses': len(self.course_info),
            'students_with_progression': len(students_with_progression)
        }

def main():
    data_path = "/Users/tikaram/Downloads/Claude Code/student-growth/growth-compass/data"
    
    analyzer = ComprehensiveFeedbackAnalyzer(data_path)
    analyzer.run_full_analysis()
    results = analyzer.generate_comprehensive_report()
    
    print(f"\\n" + "="*50)
    print("ANALYSIS COMPLETE")
    print("="*50)
    print(f"Total students identified: {results['total_students']}")
    print(f"Total courses analyzed: {results['total_courses']}")
    print(f"Students with progression data: {results['students_with_progression']}")

if __name__ == "__main__":
    main()